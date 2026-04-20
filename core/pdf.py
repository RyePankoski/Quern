import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit

w, h = letter
LEFT = 0.75 * inch
RIGHT = w - 0.75 * inch
MID = w / 2
CONTENT_W = RIGHT - LEFT


def generate_contract_pdf(data, logo_path):
    """
    Generate a contract PDF from a data dict.

    Expected keys:
        date, contract_number, seller, buyer, quality,
        quantity, uom, shipment, price, origin, grades, weights,
        governing_contract, discount, moisture, damage,
        heat_damage, foreign_materials, splits, payment,
        other_conditions, demurrage, broker
    Missing keys render as empty string.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    y = h - 0.6 * inch

    # --- LOGO ---
    logo_w = CONTENT_W
    logo_h = logo_w * (86 / 576)
    c.drawImage(logo_path, LEFT, y - logo_h, width=logo_w, height=logo_h, preserveAspectRatio=True)
    y -= logo_h + 0.15 * inch

    # --- ADDRESS ---
    c.setFont("Helvetica", 7.5)
    c.drawString(LEFT, y, "12500 Sherwood Drive, Leawood KS 66209")
    c.drawString(LEFT, y - 10, "1803 S Foothills Hwy., Suite P, Boulder, Colorado 80303")
    c.drawString(MID + inch, y, "Phone: (913) 491-3711  Email: mcdonald@mcdonaldpelz.com")
    c.drawString(MID + inch, y - 10, "Phone: (303) 543-7033  Email: bobby@mcdonaldpelz.com")
    y -= 0.4 * inch

    # --- RULE ---
    c.setStrokeColor(colors.black)
    c.setLineWidth(1.5)
    c.line(LEFT, y, RIGHT, y)
    y -= 0.35 * inch

    label_x = LEFT
    value_x = LEFT + 1.6 * inch
    line_h = 0.22 * inch

    def row(label, value):
        nonlocal y
        c.setFont("Helvetica-Bold", 9.5)
        c.drawString(label_x, y, label)
        c.setFont("Helvetica", 9.5)
        c.drawString(value_x, y, str(value) if value else '')
        y -= line_h

    # --- DATE / CONTRACT NUMBER ---
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(label_x, y, "DATE:")
    c.setFont("Helvetica", 9.5)
    c.drawString(value_x, y, str(data.get('date', '')))
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(MID + 0.5 * inch, y, "CONTRACT:")
    c.setFont("Helvetica", 9.5)
    c.drawString(MID + 1.4 * inch, y, str(data.get('contract_number', '')))
    y -= line_h

    y -= 0.04 * inch
    row("SELLER:", data.get('seller', ''))
    y -= 0.04 * inch
    row("BUYER:", data.get('buyer', ''))
    y -= 0.04 * inch
    row("QUALITY:", data.get('quality', ''))
    y -= 0.04 * inch
    row("ORIGIN:", data.get('origin', ''))
    y -= 0.04 * inch
    row("QUANTITY:", f"{data.get('quantity', '')} {data.get('uom', '')}".strip())
    row("SHIPMENT:", data.get('shipment', ''))
    row("PRICE:", data.get('price', ''))
    row("GRADES:", data.get('grades', ''))
    row("WEIGHTS:", data.get('weights', ''))

    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(label_x, y, "GOVERNING")
    y -= line_h
    row("CONTRACT:", data.get('governing_contract', ''))
    row("DISCOUNT:", data.get('discount', ''))
    row("MOISTURE:", data.get('moisture', ''))
    row("DAMAGE:", data.get('damage', ''))
    row("HEAT DAMAGE:", data.get('heat_damage', ''))

    # FOREIGN MATERIALS — wider label needs its own indent
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(label_x, y, "FOREIGN MATERIALS:")
    c.setFont("Helvetica", 9.5)
    fm_x = LEFT + 2.2 * inch
    fm_lines = simpleSplit(data.get('foreign_materials', ''), 'Helvetica', 9.5, CONTENT_W - 2.2 * inch)
    for i, line in enumerate(fm_lines):
        c.drawString(fm_x if i == 0 else value_x, y, line)
        y -= line_h

    row("SPLITS:", data.get('splits', ''))
    row("PAYMENT:", data.get('payment', ''))

    y -= 0.05 * inch
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(label_x, y, "OTHER CONDITIONS:")
    y -= line_h
    c.setFont("Helvetica", 9.5)
    other = data.get('other_conditions', '')
    if other:
        for line in simpleSplit(other, 'Helvetica', 9.5, CONTENT_W - inch):
            c.drawString(value_x, y, line)
            y -= line_h

    y -= 0.04 * inch
    row("DEMURRAGE SCHEDULE:", data.get('demurrage', ''))

    # --- BROKER ---
    y -= 0.2 * inch
    c.setFont("Helvetica", 9.5)
    c.drawString(RIGHT - 1.5 * inch, y, str(data.get('broker', '')))
    y -= line_h
    c.drawString(RIGHT - 1.0 * inch, y, "BROKER")

    # --- SELLER / BUYER SIGNATURE ---
    y -= 0.4 * inch
    c.setFont("Helvetica", 9.5)
    c.drawString(LEFT, y, "SELLER")
    c.drawString(MID + inch, y, "BUYER")

    # --- FOOTER ---
    y -= 0.3 * inch
    c.setLineWidth(0.5)
    c.line(LEFT, y, RIGHT, y)
    y -= 0.15 * inch

    footer = (
        "Please communicate any discrepancies to us within one business day of receipt of this electronic "
        "confirmation. If no discrepancies are reported it is assumed that all parties involved accept and "
        "agree to all terms as outlined above. We thank you for your business and kindly ask you to promptly "
        "sign and return a copy of this confirmation. However, the validity of this contract shall not be "
        "affected by the non-return of a signed copy."
    )
    c.setFont("Helvetica", 8)
    for line in simpleSplit(footer, 'Helvetica', 8, CONTENT_W):
        c.drawString(LEFT, y, line)
        y -= 11

    c.save()
    buffer.seek(0)
    return buffer


def generate_contract_docx(data, logo_path):
    """
    Generate a contract DOCX from the same data dict as generate_contract_pdf.
    Requires python-docx (add to requirements.txt: python-docx>=1.1.0).
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Logo
    try:
        doc.add_picture(logo_path, width=Inches(7.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass  # Logo missing — skip rather than crash

    # Address block
    addr = doc.add_paragraph()
    addr.add_run("12500 Sherwood Drive, Leawood KS 66209\n"
                 "1803 S Foothills Hwy., Suite P, Boulder, Colorado 80303\n"
                 "Phone: (913) 491-3711  |  Phone: (303) 543-7033").font.size = Pt(8)

    doc.add_paragraph("─" * 80)

    def add_row(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}  ")
        run_label.bold = True
        run_label.font.size = Pt(9.5)
        run_val = p.add_run(str(value) if value else '')
        run_val.font.size = Pt(9.5)

    # Header row
    p_header = doc.add_paragraph()
    p_header.paragraph_format.space_after = Pt(2)
    r = p_header.add_run(f"DATE:  {data.get('date', '')}          CONTRACT:  {data.get('contract_number', '')}")
    r.font.size = Pt(9.5)

    add_row("SELLER:", data.get('seller', ''))
    add_row("BUYER:", data.get('buyer', ''))
    add_row("QUALITY:", data.get('quality', ''))
    add_row("ORIGIN:", data.get('origin', ''))
    add_row("QUANTITY:", f"{data.get('quantity', '')} {data.get('uom', '')}".strip())
    add_row("SHIPMENT:", data.get('shipment', ''))
    add_row("PRICE:", data.get('price', ''))
    add_row("GRADES:", data.get('grades', ''))
    add_row("WEIGHTS:", data.get('weights', ''))
    add_row("GOVERNING CONTRACT:", data.get('governing_contract', ''))
    add_row("DISCOUNT:", data.get('discount', ''))
    add_row("MOISTURE:", data.get('moisture', ''))
    add_row("DAMAGE:", data.get('damage', ''))
    add_row("HEAT DAMAGE:", data.get('heat_damage', ''))
    add_row("FOREIGN MATERIALS:", data.get('foreign_materials', ''))
    add_row("SPLITS:", data.get('splits', ''))
    add_row("PAYMENT:", data.get('payment', ''))

    p_oc = doc.add_paragraph()
    p_oc.paragraph_format.space_after = Pt(2)
    r_oc = p_oc.add_run("OTHER CONDITIONS:\n")
    r_oc.bold = True
    r_oc.font.size = Pt(9.5)
    r_oc2 = p_oc.add_run(data.get('other_conditions', ''))
    r_oc2.font.size = Pt(9.5)

    add_row("DEMURRAGE SCHEDULE:", data.get('demurrage', ''))

    # Broker
    p_broker = doc.add_paragraph()
    p_broker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_broker.paragraph_format.space_before = Pt(12)
    r_b = p_broker.add_run(f"{data.get('broker', '')}\nBROKER")
    r_b.font.size = Pt(9.5)

    # Signatures
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(16)
    r_sig = p_sig.add_run("SELLER" + " " * 60 + "BUYER")
    r_sig.font.size = Pt(9.5)

    doc.add_paragraph("─" * 80)

    footer_text = (
        "Please communicate any discrepancies to us within one business day of receipt of this electronic "
        "confirmation. If no discrepancies are reported it is assumed that all parties involved accept and "
        "agree to all terms as outlined above. We thank you for your business and kindly ask you to promptly "
        "sign and return a copy of this confirmation. However, the validity of this contract shall not be "
        "affected by the non-return of a signed copy."
    )
    p_foot = doc.add_paragraph(footer_text)
    p_foot.runs[0].font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer